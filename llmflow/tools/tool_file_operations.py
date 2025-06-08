"""
LLMFlow - A powerful framework for building AI agents based on GAME methodology
(Goals, Actions, Memory, Environment).

File Operations Tool - Provides comprehensive file system operations with secure file handling, format conversion, and batch processing capabilities.
"""

import os
import json
import csv
import yaml
import xml.etree.ElementTree as ET
import hashlib
import mimetypes
import zipfile
import gzip
import tarfile
import shutil
import tempfile
import threading
import time
import re
import chardet
import logging
from datetime import datetime, date
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Callable, Tuple
from dataclasses import dataclass, field, asdict
from collections import defaultdict, OrderedDict
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import StringIO, BytesIO

from llmflow.tools.tool_decorator import register_tool

# Optional dependencies for document processing
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

@dataclass
class FileConfig:
    """Configuration for file operations"""
    # Security settings
    max_file_size: int = 100 * 1024 * 1024  # 100MB
    allowed_extensions: Optional[List[str]] = None
    blocked_extensions: List[str] = field(default_factory=lambda: ['.exe', '.bat', '.cmd', '.sh'])
    
    # Performance settings
    enable_cache: bool = True
    cache_size: int = 50
    parallel_workers: int = 4
    chunk_size: int = 8192
    
    # Parsing settings
    auto_detect_encoding: bool = True
    auto_detect_delimiter: bool = True
    convert_data_types: bool = True
    preserve_order: bool = False
    strip_whitespace: bool = True
    
    # General settings
    backup_enabled: bool = True
    temp_dir: Optional[str] = None
    log_level: str = 'INFO'
    encoding_fallbacks: List[str] = field(default_factory=lambda: ['utf-8', 'cp1251', 'latin1'])

class FileOperationError(Exception):
    """Base exception for file operations"""
    pass

class SecurityError(FileOperationError):
    """Security-related file operation error"""
    pass

class ValidationError(FileOperationError):
    """Validation error"""
    pass

class SecurityValidator:
    """Security validator for file operations"""
    
    def __init__(self, config: FileConfig):
        self.config = config
    
    def validate_path(self, path: str) -> bool:
        """Validate file path for security"""
        try:
            # Normalize path
            normalized = os.path.normpath(os.path.abspath(path))
            
            # Check for path traversal
            if '..' in normalized:
                return False
            
            # Check file extension
            ext = Path(path).suffix.lower()
            
            if self.config.allowed_extensions and ext not in self.config.allowed_extensions:
                return False
            
            if ext in self.config.blocked_extensions:
                return False
            
            return True
        except Exception:
            return False
    
    def validate_size(self, size: int) -> bool:
        """Validate file size"""
        return size <= self.config.max_file_size

@register_tool(tags=["file_system", "read", "file_io"])
def read_file_advanced(file_path: str, encoding: Optional[str] = None, format_hint: Optional[str] = None) -> Dict[str, Any]:
    """
    Read file content with advanced features like format detection, encoding detection, and security validation.
    
    Args:
        file_path: Path to the file to read
        encoding: Optional encoding to use. If not provided, will auto-detect.
        format_hint: Optional hint for file format
        
    Returns:
        Dictionary containing file content and metadata
    """
    if not file_path:
        return {"success": False, "error": "File path cannot be empty"}
        
    tool = FileOperationsTool(FileConfig())
    result = tool.read_file(file_path, encoding=encoding, format_hint=format_hint)
    return result.to_dict()

@register_tool(tags=["file_system", "write", "file_io"])
def write_file_advanced(file_path: str, content: Any, format_hint: Optional[str] = None, encoding: str = 'utf-8', backup: bool = True) -> Dict[str, Any]:
    """
    Write content to file with advanced features like format detection, backup and atomic writes.
    
    Args:
        file_path: Path where to write the file
        content: Content to write (can be string, dict, list, etc.)
        format_hint: Optional hint for output format
        encoding: Encoding to use (default: utf-8)
        backup: Whether to create backup of existing file
        
    Returns:
        Dictionary with operation status and metadata
    """
    if not file_path:
        return {"success": False, "error": "File path cannot be empty"}
        
    config = FileConfig()
    config.backup_enabled = backup
    tool = FileOperationsTool(config)
    result = tool.write_file(file_path, content, format_hint=format_hint, encoding=encoding)
    return result.to_dict()

@register_tool(tags=["file_system", "convert", "file_io"])
def convert_file_format(input_path: str, output_path: str, target_format: Optional[str] = None) -> Dict[str, Any]:
    """
    Convert file between formats (e.g., CSV to JSON, XML to YAML, etc.)
    
    Args:
        input_path: Source file path
        output_path: Destination file path
        target_format: Optional target format hint
        
    Returns:
        Dictionary with operation status and metadata
    """
    if not input_path or not output_path:
        return {"success": False, "error": "Input and output paths cannot be empty"}
        
    tool = FileOperationsTool(FileConfig())
    result = tool.convert_format(input_path, output_path, target_format)
    return result.to_dict()

@register_tool(tags=["file_system", "info", "file_io"])
def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Get detailed file information without reading content
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary with file metadata
    """
    if not file_path:
        return {"exists": False, "error": "File path cannot be empty"}
        
    tool = FileOperationsTool(FileConfig())
    return tool.get_file_info(file_path)

@dataclass
class FileMetadata:
    """File metadata and statistics"""
    path: str
    size: int
    created: datetime
    modified: datetime
    mime_type: str
    format: str
    encoding: Optional[str] = None
    
    # Data-specific metadata
    records_count: int = 0
    columns: List[str] = field(default_factory=list)
    column_types: Dict[str, str] = field(default_factory=dict)
    text_length: int = 0
    page_count: int = 0
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        result['created'] = self.created.isoformat()
        result['modified'] = self.modified.isoformat()
        return result

@dataclass
class OperationResult:
    """Result of file operation"""
    success: bool
    operation: str
    path: str
    data: Any = None
    metadata: Optional[FileMetadata] = None
    error: Optional[str] = None
    warnings: List[str] = field(default_factory=list)
    execution_time: float = 0.0
    
    def to_dict(self) -> Dict[str, Any]:
        result = {
            'success': self.success,
            'operation': self.operation,
            'path': self.path,
            'data': self.data,
            'error': self.error,
            'warnings': self.warnings,
            'execution_time': self.execution_time
        }
        if self.metadata:
            result['metadata'] = self.metadata.to_dict()
        return result

class FileCache:
    """Thread-safe file cache"""
    
    def __init__(self, max_size: int = 50):
        self.cache = {}
        self.access_times = {}
        self.max_size = max_size
        self.lock = threading.Lock()
    
    def get(self, key: str) -> Optional[Any]:
        with self.lock:
            if key in self.cache:
                self.access_times[key] = time.time()
                return self.cache[key]
            return None
    
    def put(self, key: str, value: Any):
        with self.lock:
            if len(self.cache) >= self.max_size:
                oldest_key = min(self.access_times.keys(), key=lambda k: self.access_times[k])
                del self.cache[oldest_key]
                del self.access_times[oldest_key]
            
            self.cache[key] = value
            self.access_times[key] = time.time()
    
    def clear(self):
        with self.lock:
            self.cache.clear()
            self.access_times.clear()

class DataTypeDetector:
    """Intelligent data type detection"""
    
    @staticmethod
    def detect_and_convert(value: str) -> Tuple[type, Any]:
        """Detect and convert data type"""
        if not value or value.strip() == '':
            return str, None
        
        value = value.strip()
        
        # Boolean
        if value.lower() in ('true', 'false', 'yes', 'no'):
            return bool, value.lower() in ('true', 'yes')
        
        # Integer
        if re.match(r'^-?\d+$', value):
            try:
                return int, int(value)
            except ValueError:
                pass
        
        # Float
        if re.match(r'^-?\d*\.?\d+([eE][+-]?\d+)?$', value):
            try:
                return float, float(value)
            except ValueError:
                pass
        
        # Date patterns
        date_patterns = [
            (r'^\d{4}-\d{2}-\d{2}$', '%Y-%m-%d'),
            (r'^\d{2}/\d{2}/\d{4}$', '%m/%d/%Y'),
            (r'^\d{2}\.\d{2}\.\d{4}$', '%d.%m.%Y'),
        ]
        
        for pattern, format_str in date_patterns:
            if re.match(pattern, value):
                try:
                    return date, datetime.strptime(value, format_str).date()
                except ValueError:
                    pass
        
        return str, value

class FileOperationsTool:
    """
    Universal File Operations Tool
    
    Handles:
    - File read/write operations
    - Structured data parsing (CSV/JSON/XML)
    - Document processing (PDF, DOCX, etc.)
    """
    
    def __init__(self, config: Optional[FileConfig] = None):
        self.config = config or FileConfig()
        self.security = SecurityValidator(self.config)
        self.cache = FileCache(self.config.cache_size) if self.config.enable_cache else None
        self.logger = self._setup_logger()
    
    def _setup_logger(self) -> logging.Logger:
        """Setup logger"""
        logger = logging.getLogger('FileOperationsTool')
        logger.setLevel(getattr(logging, self.config.log_level))
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def read_file(self, path: str, **kwargs) -> OperationResult:
        """
        Universal file reading with format auto-detection
        
        Args:
            path: File path
            **kwargs: Format-specific options
        
        Returns:
            OperationResult with parsed data
        """
        start_time = time.time()
        
        try:
            # Security validation
            if not self.security.validate_path(path):
                return OperationResult(False, 'read', path, error="Unsafe file path")
            
            if not os.path.exists(path):
                return OperationResult(False, 'read', path, error="File not found")
            
            file_size = os.path.getsize(path)
            if not self.security.validate_size(file_size):
                return OperationResult(False, 'read', path, error=f"File too large: {file_size} bytes")
            
            # Check cache
            cache_key = f"{path}:{os.path.getmtime(path)}"
            if self.cache:
                cached = self.cache.get(cache_key)
                if cached:
                    return cached
            
            # Detect format and process
            file_format = self._detect_format(path)
            data, metadata = self._read_by_format(path, file_format, **kwargs)
            
            execution_time = time.time() - start_time
            result = OperationResult(
                success=True,
                operation='read',
                path=path,
                data=data,
                metadata=metadata,
                execution_time=execution_time
            )
            
            # Cache result
            if self.cache:
                self.cache.put(cache_key, result)
            
            return result
            
        except Exception as e:
            execution_time = time.time() - start_time
            return OperationResult(
                success=False,
                operation='read',
                path=path,
                error=str(e),
                execution_time=execution_time
            )
    
    def write_file(self, path: str, data: Any, format_hint: Optional[str] = None, **kwargs) -> OperationResult:
        """
        Universal file writing with format auto-detection
        
        Args:
            path: Output file path
            data: Data to write
            format_hint: Optional format hint
            **kwargs: Format-specific options
        
        Returns:
            OperationResult with operation status
        """
        start_time = time.time()
        
        try:
            if not self.security.validate_path(path):
                return OperationResult(False, 'write', path, error="Unsafe file path")
            
            # Create directory
            os.makedirs(os.path.dirname(path), exist_ok=True)
            
            # Backup existing file
            if self.config.backup_enabled and os.path.exists(path):
                backup_path = f"{path}.backup.{int(time.time())}"
                shutil.copy2(path, backup_path)
            
            # Detect format
            file_format = format_hint or self._detect_format(path)
            
            # Atomic write
            success = self._write_by_format(path, data, file_format, **kwargs)
            
            if not success:
                return OperationResult(False, 'write', path, error="Write operation failed")
            
            # Get metadata
            metadata = self._get_file_metadata(path, file_format)
            execution_time = time.time() - start_time
            
            return OperationResult(
                success=True,
                operation='write',
                path=path,
                metadata=metadata,
                execution_time=execution_time
            )
            
        except Exception as e:
            execution_time = time.time() - start_time
            return OperationResult(
                success=False,
                operation='write',
                path=path,
                error=str(e),
                execution_time=execution_time
            )
    
    def convert_format(self, input_path: str, output_path: str, target_format: Optional[str] = None) -> OperationResult:
        """
        Convert file between formats
        
        Args:
            input_path: Source file path
            output_path: Destination file path
            target_format: Optional target format hint
            
        Returns:
            OperationResult with operation status
        """
        start_time = time.time()
        
        try:
            # Read input file
            read_result = self.read_file(input_path)
            if not read_result.success:
                return read_result
            
            # Write output file
            write_result = self.write_file(output_path, read_result.data, format_hint=target_format)
            if not write_result.success:
                return write_result
            
            execution_time = time.time() - start_time
            return OperationResult(
                success=True,
                operation='convert',
                path=output_path,
                metadata=write_result.metadata,
                execution_time=execution_time
            )
            
        except Exception as e:
            execution_time = time.time() - start_time
            return OperationResult(
                success=False,
                operation='convert',
                path=output_path,
                error=str(e),
                execution_time=execution_time
            )
    
    def _detect_format(self, path: str) -> str:
        """Auto-detect file format"""
        extension = Path(path).suffix.lower()
        
        format_map = {
            '.csv': 'csv',
            '.json': 'json',
            '.xml': 'xml',
            '.yaml': 'yaml', '.yml': 'yaml',
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.xlsx': 'excel', '.xls': 'excel',
            '.md': 'markdown',
            '.txt': 'text',
            '.log': 'text'
        }
        
        # Check if we have the required dependencies for the format
        detected_format = format_map.get(extension, 'text')
        
        if detected_format == 'pdf' and not HAS_PYPDF2:
            self.logger.warning("PyPDF2 not installed, falling back to text format")
            return 'text'
            
        if detected_format == 'docx' and not HAS_DOCX:
            self.logger.warning("python-docx not installed, falling back to text format")
            return 'text'
            
        if detected_format == 'excel' and not HAS_OPENPYXL:
            self.logger.warning("openpyxl not installed, falling back to text format")
            return 'text'
            
        if detected_format == 'markdown' and not HAS_MARKDOWN:
            self.logger.warning("markdown not installed, falling back to text format")
            return 'text'
        
        return detected_format
    
    def _read_by_format(self, path: str, file_format: str, **kwargs) -> Tuple[Any, FileMetadata]:
        """Route reading by format"""
        readers = {
            'csv': self._read_csv,
            'json': self._read_json,
            'xml': self._read_xml,
            'yaml': self._read_json,  # YAML uses JSON reader for now
            'text': self._read_text
        }
        
        # Add optional format readers if dependencies are available
        if HAS_PYPDF2:
            readers['pdf'] = self._read_pdf
        if HAS_DOCX:
            readers['docx'] = self._read_docx
        if HAS_OPENPYXL:
            readers['excel'] = self._read_excel
        if HAS_MARKDOWN:
            readers['markdown'] = self._read_markdown
        
        reader = readers.get(file_format, self._read_text)
        return reader(path, **kwargs)
    
    def _write_by_format(self, path: str, data: Any, file_format: str, **kwargs) -> bool:
        """Route writing by format"""
        if not path:
            raise ValueError("File path cannot be empty")
            
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(os.path.abspath(path)) if os.path.dirname(path) else '.', exist_ok=True)
        
        # Use temporary file for atomic write
        temp_path = f"{path}.tmp.{os.getpid()}"
        
        try:
            writers = {
                'csv': self._write_csv,
                'json': self._write_json,
                'xml': self._write_xml,
                'yaml': self._write_json,
                'markdown': self._write_text,
                'text': self._write_text
            }
            
            writer = writers.get(file_format, self._write_text)
            success = writer(temp_path, data, **kwargs)
            
            if success:
                shutil.move(temp_path, path)
                return True
            else:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                return False
                
        except Exception:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            raise
    
    def _get_file_metadata(self, path: str, file_format: str) -> FileMetadata:
        """Get comprehensive file metadata"""
        file_path = Path(path)
        stat = file_path.stat()
        
        return FileMetadata(
            path=str(file_path),
            size=stat.st_size,
            created=datetime.fromtimestamp(stat.st_ctime),
            modified=datetime.fromtimestamp(stat.st_mtime),
            mime_type=mimetypes.guess_type(str(file_path))[0] or 'application/octet-stream',
            format=file_format
        )
    
    def _read_text(self, path: str, **kwargs) -> Tuple[str, FileMetadata]:
        """Read plain text file"""
        encoding = kwargs.get('encoding', 'utf-8')
        
        with open(path, 'r', encoding=encoding) as f:
            content = f.read()
        
        metadata = self._get_file_metadata(path, 'text')
        metadata.text_length = len(content)
        metadata.encoding = encoding
        
        return content, metadata
    
    def _write_text(self, path: str, data: Any, **kwargs) -> bool:
        """Write text file"""
        encoding = kwargs.get('encoding', 'utf-8')
        
        content = data
        if isinstance(data, dict) and 'markdown' in data:
            content = data['markdown']
        elif not isinstance(data, str):
            content = str(data)
        
        with open(path, 'w', encoding=encoding) as f:
            f.write(content)
        
        return True
    
    def _read_json(self, path: str, **kwargs) -> Tuple[Any, FileMetadata]:
        """Read and parse JSON file"""
        encoding = kwargs.get('encoding', 'utf-8')
        
        with open(path, 'r', encoding=encoding) as f:
            content = f.read()
        
        # Remove comments if needed
        if '//' in content or '/*' in content:
            content = re.sub(r'//.*', '', content)
            content = re.sub(r'/\*.*?\*/', '', content, flags=re.DOTALL)
        
        if self.config.preserve_order:
            data = json.loads(content, object_pairs_hook=OrderedDict)
        else:
            data = json.loads(content)
        
        # Generate metadata
        metadata = self._get_file_metadata(path, 'json')
        metadata.encoding = encoding
        
        if isinstance(data, list):
            metadata.records_count = len(data)
            if data and isinstance(data[0], dict):
                metadata.columns = list(data[0].keys())
        elif isinstance(data, dict):
            metadata.records_count = 1
            metadata.columns = list(data.keys())
        
        return data, metadata
    
    def _write_json(self, path: str, data: Any, **kwargs) -> bool:
        """Write JSON file"""
        encoding = kwargs.get('encoding', 'utf-8')
        indent = kwargs.get('indent', 2)
        ensure_ascii = kwargs.get('ensure_ascii', False)
        
        with open(path, 'w', encoding=encoding) as f:
            json.dump(data, f, indent=indent, ensure_ascii=ensure_ascii, default=str)
        
        return True
    
    def _read_csv(self, path: str, **kwargs) -> Tuple[List[Dict], FileMetadata]:
        """Read and parse CSV file"""
        encoding = kwargs.get('encoding', 'utf-8')
        delimiter = kwargs.get('delimiter', ',')
        
        data = []
        with open(path, 'r', encoding=encoding) as f:
            reader = csv.DictReader(f, delimiter=delimiter)
            headers = reader.fieldnames or []
            
            for row in reader:
                if self.config.strip_whitespace:
                    row = {k: v.strip() if isinstance(v, str) else v for k, v in row.items()}
                
                # Type conversion
                if self.config.convert_data_types:
                    for key, value in row.items():
                        if isinstance(value, str):
                            _, converted = DataTypeDetector.detect_and_convert(value)
                            row[key] = converted
                
                data.append(row)
        
        # Generate metadata
        metadata = self._get_file_metadata(path, 'csv')
        metadata.records_count = len(data)
        metadata.columns = headers
        metadata.encoding = encoding
        
        if data:
            metadata.column_types = {k: type(v).__name__ for k, v in data[0].items()}
        
        return data, metadata
    
    def _write_csv(self, path: str, data: List[Dict], **kwargs) -> bool:
        """Write CSV file"""
        if not data:
            return False
        
        encoding = kwargs.get('encoding', 'utf-8')
        delimiter = kwargs.get('delimiter', ',')
        headers = kwargs.get('headers') or list(data[0].keys())
        
        with open(path, 'w', encoding=encoding, newline='') as f:
            writer = csv.DictWriter(f, fieldnames=headers, delimiter=delimiter)
            writer.writeheader()
            writer.writerows(data)
        
        return True
    
    def _read_xml(self, path: str, **kwargs) -> Tuple[Dict, FileMetadata]:
        """Read and parse XML file"""
        tree = ET.parse(path)
        root = tree.getroot()
        data = self._xml_to_dict(root)
        
        metadata = self._get_file_metadata(path, 'xml')
        metadata.records_count = 1
        
        return data, metadata
    
    def _write_xml(self, path: str, data: Dict, **kwargs) -> bool:
        """Write XML file"""
        encoding = kwargs.get('encoding', 'utf-8')
        root_name = kwargs.get('root_name', 'root')
        
        root = ET.Element(root_name)
        self._dict_to_xml(data, root)
        
        tree = ET.ElementTree(root)
        tree.write(path, encoding=encoding, xml_declaration=True)
        
        return True
    
    def _xml_to_dict(self, element: ET.Element) -> Dict[str, Any]:
        """Convert XML element to dictionary"""
        result = {}
        
        # Handle attributes
        if element.attrib:
            result['@attributes'] = element.attrib
        
        # Handle text
        if element.text and element.text.strip():
            text = element.text.strip()
            if self.config.convert_data_types:
                _, text = DataTypeDetector.detect_and_convert(text)
            
            if len(element) == 0:
                return text
            else:
                result['#text'] = text
        
        # Handle children
        children = {}
        for child in element:
            child_data = self._xml_to_dict(child)
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag in children:
                if not isinstance(children[tag], list):
                    children[tag] = [children[tag]]
                children[tag].append(child_data)
            else:
                children[tag] = child_data
        
        result.update(children)
        return result
    
    def _dict_to_xml(self, data: Dict, parent: ET.Element):
        """Convert dictionary to XML elements"""
        for key, value in data.items():
            if key == '@attributes':
                parent.attrib.update(value)
            elif key == '#text':
                parent.text = str(value)
            elif isinstance(value, list):
                for item in value:
                    elem = ET.SubElement(parent, key)
                    if isinstance(item, dict):
                        self._dict_to_xml(item, elem)
                    else:
                        elem.text = str(item)
            elif isinstance(value, dict):
                elem = ET.SubElement(parent, key)
                self._dict_to_xml(value, elem)
            else:
                elem = ET.SubElement(parent, key)
                elem.text = str(value)

    def get_file_info(self, path: str) -> Dict[str, Any]:
        """Get file information without reading content"""
        if not path:
            return {'exists': False, 'error': 'File path cannot be empty'}
            
        if not os.path.exists(path):
            return {'exists': False, 'error': 'File not found'}
        
        try:
            file_format = self._detect_format(path)
            metadata = self._get_file_metadata(path, file_format)
            
            return {
                'exists': True,
                'format': file_format,
                'size': metadata.size,
                'created': metadata.created.isoformat(),
                'modified': metadata.modified.isoformat(),
                'mime_type': metadata.mime_type
            }
        except Exception as e:
            return {'exists': True, 'error': str(e)}

    def _read_pdf(self, path: str, **kwargs) -> Tuple[str, FileMetadata]:
        """Read PDF document"""
        if not HAS_PYPDF2:
            raise ImportError("PyPDF2 not installed")
        
        text = ""
        page_count = 0
        
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        metadata = self._get_file_metadata(path, 'pdf')
        metadata.text_length = len(text)
        metadata.page_count = page_count
        
        return text, metadata
    
    def _read_docx(self, path: str, **kwargs) -> Tuple[str, FileMetadata]:
        """Read DOCX document"""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed")
        
        doc = DocxDocument(path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        metadata = self._get_file_metadata(path, 'docx')
        metadata.text_length = len(text)
        metadata.page_count = len(doc.paragraphs)  # Approximate
        
        return text, metadata
    
    def _read_excel(self, path: str, **kwargs) -> Tuple[Dict[str, List[Dict]], FileMetadata]:
        """Read Excel file"""
        if not HAS_OPENPYXL:
            raise ImportError("openpyxl not installed")
        
        workbook = openpyxl.load_workbook(path, data_only=True)
        data = {}
        total_records = 0
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            rows = list(sheet.iter_rows(values_only=True))
            
            if not rows:
                continue
            
            headers = [str(cell) if cell is not None else f"Column_{i}" for i, cell in enumerate(rows[0])]
            sheet_data = []
            
            for row in rows[1:]:
                record = {}
                for i, cell in enumerate(row):
                    if i < len(headers):
                        value = cell
                        if self.config.convert_data_types and isinstance(value, str):
                            _, value = DataTypeDetector.detect_and_convert(value)
                        record[headers[i]] = value
                sheet_data.append(record)
            
            data[sheet_name] = sheet_data
            total_records += len(sheet_data)
        
        metadata = self._get_file_metadata(path, 'excel')
        metadata.records_count = total_records
        
        return data, metadata
    
    def _read_markdown(self, path: str, **kwargs) -> Tuple[str, FileMetadata]:
        """Read Markdown file"""
        if not HAS_MARKDOWN:
            raise ImportError("markdown not installed")
            
        encoding = kwargs.get('encoding', 'utf-8')
        
        with open(path, 'r', encoding=encoding) as f:
            content = f.read()
        
        # Convert to HTML if markdown library available
        html_content = markdown.markdown(content)
        content = {'markdown': content, 'html': html_content}
        
        metadata = self._get_file_metadata(path, 'markdown')
        metadata.text_length = len(content['markdown'])
        metadata.encoding = encoding
        
        return content, metadata 